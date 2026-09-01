"""Unit tests for WordParser -- section-aware Word text chunking (T026).

The WordParser is a pure text parser: it consumes the text produced by
text_extractor.extract_text (which preserves heading markers, list markers,
and table rows) and chunks it like markdown.  These tests mirror the
test_pdf_parser pattern: a fixture runs extract_text on a sample docx and
hands the resulting text to the parser.
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from rag_mcp.parsers.text_extractor import extract_text, TextExtractionError
from rag_mcp.parsers.word_parser import WordParser

FIXTURES = Path(__file__).parents[2] / "fixtures" / "samples"


@pytest.fixture
def parser() -> WordParser:
    return WordParser()


@pytest.fixture
def design_text() -> str:
    """extract_text output for the multi-level-heading sample docx."""
    return extract_text((FIXTURES / "design.docx").read_bytes(), "word")


def _png_bytes(w: int = 2, h: int = 2) -> bytes:
    """Build a minimal valid PNG using integer byte lists (no escapes)."""
    sig = bytes([137, 80, 78, 71, 13, 10, 26, 10])
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    pixel = bytes([255, 0, 0])
    raw = b""
    for _ in range(h):
        raw += bytes([0]) + pixel * w
    idat = zlib.compress(raw)

    def _chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


def _docx_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# 1. Heading extraction
# --------------------------------------------------------------------------- #

class TestWordParserHeadings:
    def test_extracts_all_headings(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        headings = [c for c in chunks if c["chunk_type"] == "heading"]
        assert headings, "expected heading chunks"

        titles = " ".join(c["content_text"] for c in headings)
        assert "Architecture Design" in titles, titles
        assert "System Overview" in titles, titles
        assert "Data Flow" in titles, titles

    def test_heading_hierarchy_levels(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        paths = [c["section_path"] for c in chunks]
        assert "# Architecture Design" in paths, paths
        assert "# Architecture Design > ## System Overview" in paths, paths
        assert "# Architecture Design > ## System Overview > ### Data Flow" in paths, paths
        assert "# Architecture Design > ## API Design" in paths, paths
        assert "# Architecture Design > ## Deployment" in paths, paths


# --------------------------------------------------------------------------- #
# 2. Paragraph extraction
# --------------------------------------------------------------------------- #

class TestWordParserParagraphs:
    def test_extracts_body_paragraphs(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        paras = [c for c in chunks if c["chunk_type"] == "paragraph"]
        assert len(paras) >= 4, f"expected >=4 paragraphs, got {len(paras)}"

        text = " ".join(c["content_text"] for c in paras)
        assert "front-end/back-end separation" in text, text
        assert "API gateway" in text, text
        assert "OpenAPI 3.0" in text, text
        assert "Docker" in text, text


# --------------------------------------------------------------------------- #
# 3. Table extraction
# --------------------------------------------------------------------------- #

class TestWordParserTables:
    def test_extracts_table_rows(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        tables = [c for c in chunks if c["chunk_type"] == "table"]
        assert len(tables) == 1, f"expected 1 table chunk, got {len(tables)}"

        content = tables[0]["content_text"]
        assert "Field" in content and "Type" in content, content
        assert "id" in content and "integer" in content, content
        assert "name" in content and "string" in content, content


# --------------------------------------------------------------------------- #
# 4. section_path format
# --------------------------------------------------------------------------- #

class TestWordParserSectionPath:
    def test_section_path_hash_format(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        for c in chunks:
            path = c["section_path"]
            if not path:
                continue
            for part in path.split(" > "):
                assert part.startswith("#"), f"path part {part!r} must start with '#'"
                assert part[1] in ("#", " "), f"malformed path part {part!r}"


# --------------------------------------------------------------------------- #
# 5. parent_section_path tracking
# --------------------------------------------------------------------------- #

class TestWordParserParentTracking:
    def test_parent_section_path_correct(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        headings = {h["section_path"]: h for h in chunks if h["chunk_type"] == "heading"}
        assert (
            headings["# Architecture Design > ## System Overview > ### Data Flow"]["parent_section_path"]
            == "# Architecture Design > ## System Overview"
        )
        assert (
            headings["# Architecture Design > ## System Overview"]["parent_section_path"]
            == "# Architecture Design"
        )
        assert headings["# Architecture Design"]["parent_section_path"] == ""

    def test_child_section_path_prefixed_by_parent(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        children = [c for c in chunks if c["parent_section_path"]]
        assert children, "expected chunks with a non-empty parent_section_path"
        for c in children:
            parent = c["parent_section_path"]
            section = c["section_path"]
            assert section.startswith(parent + " > "), f"{section!r} should start with {parent!r} + ' > '"


# --------------------------------------------------------------------------- #
# 6. chunk_type values
# --------------------------------------------------------------------------- #

class TestWordParserChunkTypes:
    def test_chunk_types_valid(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        valid = {"heading", "paragraph", "list", "table"}
        present = {c["chunk_type"] for c in chunks}
        assert present <= valid, f"unexpected chunk types: {present - valid}"
        assert "heading" in present
        assert "paragraph" in present
        assert "table" in present


# --------------------------------------------------------------------------- #
# 7. List items
# --------------------------------------------------------------------------- #

class TestWordParserLists:
    def test_extracts_list_items(self, parser: WordParser) -> None:
        doc = Document()
        doc.add_heading("Lists Section", level=1)
        doc.add_paragraph("intro prose")
        doc.add_paragraph("first bullet", style="List Bullet")
        doc.add_paragraph("second bullet", style="List Bullet")
        doc.add_paragraph("numbered item", style="List Number")
        text = extract_text(_docx_to_bytes(doc), "word")

        chunks = parser.parse(text, "lists.docx")
        lists = [c for c in chunks if c["chunk_type"] == "list"]
        assert lists, "expected at least one list chunk"

        joined = " ".join(c["content_text"] for c in lists)
        assert "first bullet" in joined, joined
        assert "second bullet" in joined, joined
        assert "numbered item" in joined, joined


# --------------------------------------------------------------------------- #
# 8. Empty / edge input
# --------------------------------------------------------------------------- #

class TestWordParserEmpty:
    def test_empty_text_returns_empty(self, parser: WordParser) -> None:
        assert parser.parse("") == []
        assert parser.parse("   ") == []
        assert parser.parse(chr(10) * 4) == []

    def test_empty_docx_raises_at_extraction(self) -> None:
        """An empty docx is rejected at the extraction stage (FR-019), like scanned PDFs."""
        with pytest.raises(TextExtractionError):
            extract_text((FIXTURES / "empty.docx").read_bytes(), "word")


# --------------------------------------------------------------------------- #
# 9. Embedded objects skipped
# --------------------------------------------------------------------------- #

class TestWordParserEmbeddedObjects:
    def test_embedded_image_skipped(self, parser: WordParser) -> None:
        """An image paragraph yields no text -> the extractor skips it; the
        parser then never sees image bytes, only the surrounding prose."""
        doc = Document()
        doc.add_heading("With Image", level=1)
        doc.add_paragraph("text before image")
        doc.add_picture(io.BytesIO(_png_bytes(2, 2)))
        doc.add_paragraph("text after image")
        text = extract_text(_docx_to_bytes(doc), "word")

        chunks = parser.parse(text, "with_image.docx")
        paras = [c for c in chunks if c["chunk_type"] == "paragraph"]
        para_text = " ".join(c["content_text"] for c in paras)
        assert "text before image" in para_text, para_text
        assert "text after image" in para_text, para_text
        assert all("PNG" not in c["content_text"] for c in chunks)


# --------------------------------------------------------------------------- #
# 10. No-heading fallback
# --------------------------------------------------------------------------- #

class TestWordParserNoHeadings:
    def test_synthesizes_root_heading(self, parser: WordParser) -> None:
        """A heading-less doc synthesizes '# {filename}' as section_path."""
        doc = Document()
        doc.add_paragraph("just a paragraph of prose")
        doc.add_paragraph("another paragraph")
        text = extract_text(_docx_to_bytes(doc), "word")

        chunks = parser.parse(text, "plain.docx")
        assert chunks, "expected chunks for no-heading document"
        for c in chunks:
            assert c["section_path"] == "# plain.docx", c["section_path"]
            assert c["parent_section_path"] == "", c["parent_section_path"]


# --------------------------------------------------------------------------- #
# 11. Required chunk fields & line numbers
# --------------------------------------------------------------------------- #

class TestWordParserRequiredFields:
    def test_all_chunks_have_required_fields(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        assert chunks, "expected chunks from design.docx"

        required = {
            "content_text",
            "section_path",
            "start_line",
            "end_line",
            "parent_section_path",
            "token_count",
            "chunk_type",
        }
        for i, c in enumerate(chunks):
            missing = required - set(c.keys())
            assert not missing, f"chunk {i} missing keys: {missing}"

            assert isinstance(c["content_text"], str) and c["content_text"].strip()
            assert isinstance(c["section_path"], str)
            assert isinstance(c["start_line"], int) and c["start_line"] >= 1
            assert isinstance(c["end_line"], int) and c["end_line"] >= c["start_line"]
            assert isinstance(c["parent_section_path"], str)
            assert isinstance(c["token_count"], int) and c["token_count"] > 0
            assert isinstance(c["chunk_type"], str)

    def test_line_numbers_within_bounds(self, parser: WordParser, design_text: str) -> None:
        chunks = parser.parse(design_text, "design.docx")
        n = len(design_text.splitlines())
        for c in chunks:
            assert 1 <= c["start_line"] <= n, c
            assert 1 <= c["end_line"] <= n, c
            assert c["end_line"] >= c["start_line"]


# --------------------------------------------------------------------------- #
# 12. Corrupt binary input
# --------------------------------------------------------------------------- #

class TestWordParserCorruptBinary:
    def test_corrupt_docx_raises_at_extraction(self) -> None:
        # A corrupt docx is rejected at the extraction stage (FR-019), like
        # empty/scanned binaries; the parser never sees it, so no chunks are
        # produced. An unfinished version therefore cannot participate in
        # retrieval and the previous version stays authoritative.
        with pytest.raises(TextExtractionError):
            extract_text((FIXTURES / "corrupt.docx").read_bytes(), "word")
