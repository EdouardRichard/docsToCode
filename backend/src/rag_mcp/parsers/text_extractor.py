"""
Binary format text extraction dispatcher (003, FR-011).

Provides extract_text(raw_bytes, fmt) -> str for Word and PDF formats.
Text formats bypass this module -- their raw content is already text.

Extraction failures raise an exception with a clear cause (FR-019).
"""

from __future__ import annotations

import logging

logger = logging.getLogger(__name__)

BINARY_FORMATS = frozenset({"word", "pdf"})


class TextExtractionError(Exception):
    """Raised when text extraction fails (FR-011/FR-019)."""


def extract_text(raw_bytes, fmt):
    """Extract text from binary format raw bytes."""
    if fmt == "word":
        return _extract_word_text(raw_bytes)
    elif fmt == "pdf":
        return _extract_pdf_text(raw_bytes)
    raise ValueError(f"Format {fmt!r} is not a binary format")


def _extract_word_text(raw_bytes):
    """Extract text from Word .docx with heading structure preserved (FR-011)."""
    import io
    try:
        from docx import Document
    except ImportError as exc:
        raise TextExtractionError("python-docx not installed") from exc
    try:
        doc = Document(io.BytesIO(raw_bytes))
    except Exception as exc:
        raise TextExtractionError(f"Failed to open Word doc: {exc}") from exc
    parts = []
    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
        style_name = ""
        try:
            style_name = para.style.name or ""
        except Exception:
            pass
        # Preserve heading structure with markdown-style markers
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip() or "1")
            except ValueError:
                level = 1
            parts.append("#" * level + " " + text)
        elif style_name == "Title":
            parts.append("# " + text)
        elif style_name.startswith("List Bullet") or style_name.startswith("List Number"):
            # Preserve list structure with markdown-style "- " markers so the
            # downstream parser can emit chunk_type="list" blocks.
            parts.append("- " + text)
        else:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells: parts.append("| " + " | ".join(cells) + " |")
    result = chr(10).join(parts)
    if not result.strip():
        raise TextExtractionError("Word document is empty (FR-019)")
    return result


# --- Column-aware PDF extraction (003 FR-006 / Clarification Q2) ---------
#
# Reading-order strategy (research.md section 1.6):
#   1. Extract word coordinates per page.
#   2. Detect full-height column gutters from the union of word x-intervals.
#   3. Multi-column pages are rebuilt column by column (left column top to
#      bottom, then the next column) so text enters chunks in reading order
#      instead of linearly interleaving the columns.
#   4. When no clean gutter exists but per-line gap analysis suggests an
#      (undetectable) multi-column layout -- e.g. full-width headings
#      spanning the gutter -- the page degrades to linear extraction and the
#      reason is annotated in the page marker (edge case: column detection
#      failure degrades to linear with annotation). A page is never
#      discarded because column detection failed.

_MIN_GUTTER_RATIO = 0.03       # gutter >= 3% of page width counts as a boundary
_MIN_GUTTER_ABS = 12.0         # ...and at least 12pt (absolute floor)
_MIN_WORDS_PER_COLUMN = 2      # fewer words in a region -> not a real column
_LINE_TOP_TOL = 3.0            # tolerance when grouping words into visual lines
_GAP_CLUSTER_TOL = 15.0        # gaps within 15pt count as the same gutter
_SUSPICIOUS_LINE_RATIO = 0.25  # >=25% bimodal lines -> suspicious multi-column


def _min_gutter_width(page_width: float) -> float:
    return max(_MIN_GUTTER_ABS, page_width * _MIN_GUTTER_RATIO)


def _column_of(word: dict, boundaries) -> int:
    """Index of the column a word belongs to (by horizontal center)."""
    center = (word["x0"] + word["x1"]) / 2.0
    for i, boundary in enumerate(boundaries):
        if center < boundary:
            return i
    return len(boundaries)


def _detect_columns(words, page_width: float):
    """Detect column boundaries via full-height x-coverage gutters.

    Returns "(column_count, boundaries)"; "(1, [])" when no clean gutter
    exists (single column, or a multi-column layout this detector cannot
    resolve -- see _suspicious_multicolumn_reason).
    """
    if len(words) < 2 * _MIN_WORDS_PER_COLUMN:
        return 1, []
    gutter = _min_gutter_width(page_width)
    intervals = sorted((w["x0"], w["x1"]) for w in words)
    merged = [list(intervals[0])]
    for x0, x1 in intervals[1:]:
        if x0 <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], x1)
        else:
            merged.append([x0, x1])
    gaps = [
        (a[1], b[0])
        for a, b in zip(merged, merged[1:])
        if b[0] - a[1] >= gutter
    ]
    if not gaps:
        return 1, []
    boundaries = [(g0 + g1) / 2.0 for g0, g1 in gaps]
    counts = [0] * (len(boundaries) + 1)
    for w in words:
        counts[_column_of(w, boundaries)] += 1
    if any(c < _MIN_WORDS_PER_COLUMN for c in counts):
        # A stray word inside whitespace must not fabricate a column.
        return 1, []
    return len(boundaries) + 1, boundaries


def _group_into_lines(words):
    """Group words into visual lines by top coordinate (anchor-based)."""
    lines = []
    for w in sorted(words, key=lambda w: (w["top"], w["x0"])):
        if lines and abs(w["top"] - lines[-1][0]) <= _LINE_TOP_TOL:
            lines[-1][1].append(w)
        else:
            lines.append([w["top"], [w]])
    return [line_words for _, line_words in lines]


def _reconstruct_columnar(words, boundaries) -> str:
    """Rebuild page text in reading order: column by column (left to
    right), each column top-to-bottom (FR-006)."""
    columns = [[] for _ in range(len(boundaries) + 1)]
    for w in words:
        columns[_column_of(w, boundaries)].append(w)
    out_lines = []
    for column_words in columns:
        for line_words in _group_into_lines(column_words):
            line_words = sorted(line_words, key=lambda w: w["x0"])
            out_lines.append(" ".join(w["text"] for w in line_words))
    return chr(10).join(out_lines)


def _suspicious_multicolumn_reason(words, page_width: float):
    """Degradation heuristic for an undetectable multi-column layout.

    Returns a short reason string when the page looks multi-column (most
    lines share a consistent intra-line x-gap) yet no clean full-height
    gutter exists -- e.g. full-width headings span the gutter -- and None
    otherwise (single column, or a cleanly detectable multi-column page).
    """
    n_cols, _ = _detect_columns(words, page_width)
    if n_cols >= 2:
        return None  # cleanly detectable: not a degradation case
    if len(words) < 8:
        return None
    lines = _group_into_lines(words)
    if len(lines) < 4:
        return None
    gutter = _min_gutter_width(page_width)
    gap_positions = []
    for line_words in lines:
        line_words = sorted(line_words, key=lambda w: w["x0"])
        for a, b in zip(line_words, line_words[1:]):
            if b["x0"] - a["x1"] >= gutter:
                gap_positions.append((a["x1"] + b["x0"]) / 2.0)
    if not gap_positions:
        return None
    gap_positions.sort()
    clusters = [[gap_positions[0]]]
    for pos in gap_positions[1:]:
        if pos - clusters[-1][-1] <= _GAP_CLUSTER_TOL:
            clusters[-1].append(pos)
        else:
            clusters.append([pos])
    largest = max(clusters, key=len)
    if len(largest) >= max(2, int(_SUSPICIOUS_LINE_RATIO * len(lines))):
        return "column-layout-inconclusive"
    return None


def _extract_page_text(page, words):
    """Extract one page's text column-aware (FR-006).

    Returns "(text, annotation)": annotation is None for clean single- or
    multi-column pages, or a "linear-fallback: <reason>" string when the
    layout looks multi-column but no clean gutter could be detected.
    """
    n_cols, boundaries = _detect_columns(words, page.width)
    if n_cols >= 2:
        return _reconstruct_columnar(words, boundaries), None
    reason = _suspicious_multicolumn_reason(words, page.width)
    if reason:
        return (page.extract_text() or ""), f"linear-fallback: {reason}"
    return (page.extract_text() or ""), None


def _extract_pdf_text(raw_bytes):
    """Extract text from PDF with page markers preserved (FR-011) and
    multi-column reading order retained (FR-006 / Clarification Q2)."""
    import io
    try:
        import pdfplumber
    except ImportError as exc:
        raise TextExtractionError("pdfplumber not installed") from exc
    try:
        pages_text = []
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            if len(pdf.pages) == 0:
                raise TextExtractionError("PDF has no pages (FR-019)")
            for i, page in enumerate(pdf.pages):
                try:
                    words = page.extract_words() or []
                except Exception:  # noqa: BLE001 - layout analysis is best-effort
                    words = []
                page_text, annotation = _extract_page_text(page, words)
                marker = f"=== PAGE {i+1} ==="
                if annotation:
                    marker = f"=== PAGE {i+1} ({annotation}) ==="
                    logger.warning(
                        "PDF page %d: column detection degraded (%s); "
                        "falling back to linear extraction",
                        i + 1, annotation,
                    )
                # Preserve page boundaries with markers for the parser
                pages_text.append(marker + chr(10) + page_text)
    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError(f"Failed to open PDF: {exc}") from exc
    result = (chr(10) + chr(10)).join(pages_text)
    # Check for actual text content (not just page markers)
    actual_text = chr(10).join(
        line for line in result.splitlines()
        if line.strip() and not line.strip().startswith("=== PAGE")
    )
    if not actual_text.strip():
        raise TextExtractionError("No text layer in PDF -- scanned? (FR-019)")
    return result
